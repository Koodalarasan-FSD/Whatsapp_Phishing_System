from django.shortcuts import render,redirect
from .models import DateTimeRecord
from datetime import datetime
from django.http import FileResponse # Used for download the file

# Install Packages - pip install document, docx, python-docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

import os

# Create your views here.
def index(request,id):
    # Check if a record with the given id already exists
    existing_records = DateTimeRecord.objects.filter(id=id)
    
    if existing_records.exists():
        # If there are existing records, update failed_click_datetime
        current_record = existing_records.first()
        current_record.failed_click_datetime = datetime.now()
        current_record.save()
    else:
        return render(request,'404.html')
    
    # Pass the id as context data to the template
    context={'id':id}
        
    return render(request,'index.html',context)

def login_submit(request):
    if request.method=="POST":
        
        # Extract id from form which is hidden
        id=request.POST.get('id')

        # Check if a record with the given id already exists
        existing_records = DateTimeRecord.objects.filter(id=id)
    
        if existing_records.exists():
            # If there are existing records, update failed_responded_datetime field
            current_record = existing_records.first()
            current_record.failed_responded_datetime = datetime.now()
            current_record.save()
        else:
            return render(request,'404.html')
    
    return render(request,'404.html')

# URL to render admin_report page
def admin_report(request):
    return render(request,'admin_report.html')

# Below both functions set_cell_format and set_cell_margins for table content in document(report file)
def set_cell_format(cell,is_header=False):
    run = cell.paragraphs[0].runs[0]
    run.font.bold = is_header
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(10)
def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    cell.paragraphs[0].paragraph_format.space_before = Pt(top)
    cell.paragraphs[0].paragraph_format.space_after = Pt(bottom)
    cell.paragraphs[0].paragraph_format.left_indent = Pt(left)
    cell.paragraphs[0].paragraph_format.right_indent = Pt(right)

# Function to generate and download admin report file
def generate_download_report(request):
    # Retrieve all objects from the table
    objects = DateTimeRecord.objects.all()

    # Check if there are any records
    if not objects.exists():
        # Handle case where there are no records
        return redirect('admin_report')

    # Create a new Word document
    doc=Document()

    # Add a table to the document
    fields = [field.name for field in DateTimeRecord._meta.get_fields()]
    table = doc.add_table(rows=1, cols=len(fields), style='Table Grid')

    # Add the header row to the table
    header_row = table.rows[0]
    for col_num, field_name in enumerate(fields):
        cell = header_row.cells[col_num]
        cell.text = field_name
        set_cell_format(cell, is_header=True)  # Apply formatting to the header cell
        set_cell_margins(cell, top=5, bottom=5, left=5, right=5)  # Adjust spacing for the header row

    # Iterate through data and write to the document
    for record in objects:
        row_cells = table.add_row().cells
        for col_num, field_name in enumerate(fields):
            value = getattr(record, field_name)
            cell = row_cells[col_num]
            cell.text = str(value)
            set_cell_format(cell)  # Apply formatting to the data cell
            set_cell_margins(cell, top=5, bottom=5, left=5, right=5)  # Adjust spacing for the data rows

    # Specify the full path for saving the document
    output_path = os.path.join(os.getcwd(), 'output_report.docx')

    # Save the document
    doc.save(output_path)

    # Use FileResponse to send(downloads) the file 
    response=FileResponse(open(output_path,'rb'),as_attachment=True)
    return response
    